import asyncio
import gc
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from telegram import Document, PhotoSize, Update
from telegram.ext import ContextTypes

from src.config import settings
from src.database import db
from src.logger import logger
from src.ollama_client import llm
from src.utils.tts_manager import convert_to_ogg, text_to_speech

INBOX = Path("/data/obsidian_vault/00-Inbox")
ATTACHMENTS = Path("/data/obsidian_vault/Attachments")
MAX_FILE_SIZE = 50 * 1024 * 1024
TEMP_DIR = Path("/tmp/rafita_uploads")

CATEGORY_KEYWORDS = {
    "proyecto": "01-Proyectos",
    "proyectos": "01-Proyectos",
    "trabajo": "02-Areas/Trabajo",
    "laboral": "02-Areas/Trabajo",
    "finanza": "02-Areas/Finanzas",
    "financiero": "02-Areas/Finanzas",
    "factura": "02-Areas/Finanzas",
    "recibo": "02-Areas/Finanzas",
    "salud": "02-Areas/Salud",
    "medico": "02-Areas/Salud",
    "casa": "02-Areas/Casa",
    "hogar": "02-Areas/Casa",
    "recurso": "03-Recursos",
    "guia": "03-Recursos",
    "manual": "03-Recursos",
    "tutorial": "03-Recursos",
    "archivo": "04-Archivo",
    "historia": "04-Archivo",
    "inbox": "00-Inbox",
    "zettle": "05-Zettelkasten",
    "diario": "06-Diario",
    "nota atomica": "05-Zettelkasten",
}

EXTENSION_FOLDERS = {
    ".pdf": "03-Recursos",
    ".doc": "03-Recursos",
    ".docx": "03-Recursos",
    ".xls": "02-Areas/Finanzas",
    ".xlsx": "02-Areas/Finanzas",
    ".csv": "02-Areas/Finanzas",
    ".jpg": "Attachments",
    ".jpeg": "Attachments",
    ".png": "Attachments",
    ".gif": "Attachments",
    ".mp4": "Attachments",
    ".mov": "Attachments",
    ".mp3": "Attachments",
    ".ogg": "Attachments",
    ".wav": "Attachments",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp"}

EXTRACTABLE_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".docx"}


def _safe_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r'[<>:"/\\|?*]', "_", name)
    name = re.sub(r"\s+", "_", name)
    return name[:100]


def _safe_vault_subpath(vault_root: Path, subpath: str) -> Path:
    """Sanitize a subpath within the vault, preventing path traversal."""
    subpath = subpath.strip()
    subpath = subpath.replace("\\", "/")
    parts = [p for p in subpath.split("/") if p and p not in (".", "..")]
    result = vault_root.joinpath(*parts)
    if not str(result.resolve()).startswith(str(vault_root.resolve())):
        raise ValueError("Path traversal blocked: %s" % subpath)
    return result


def _extract_text_from_file(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n\n".join(parts)
        elif ext == ".docx":
            from docx import Document

            doc = Document(str(file_path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(cells))
            return "\n".join(parts)
        elif ext in (".txt", ".md", ".csv"):
            return file_path.read_text(encoding="utf-8", errors="replace")
        else:
            return ""
    except Exception as e:
        logger.warning("Text extraction failed for %s: %s", file_path.name, e)
        return ""


def _generate_frontmatter(
    file_path: Path, note_type: str, tags: list, extracted_summary: str
) -> str:
    now = datetime.now()
    note_id = now.strftime("%Y%m%d-%H%M")
    title = file_path.stem.replace("_", " ")
    tags_yaml = "[%s]" % ", ".join(tags) if tags else "[]"
    fm = (
        "---\n"
        "id: %s\n"
        'title: "%s"\n'
        "type: %s\n"
        "tags: %s\n"
        "status: abierto\n"
        "created: %s\n"
        "updated: %s\n"
        "related: []\n"
        "source_file: %s\n"
        "---\n"
    ) % (
        note_id,
        title,
        note_type,
        tags_yaml,
        now.strftime("%Y-%m-%d"),
        now.strftime("%Y-%m-%d"),
        file_path.name,
    )
    return fm


def _create_companion_note(
    vault_path: Path,
    saved_file: Path,
    extracted_text: str,
    note_type: str,
    tags: list,
    summary: str,
) -> Path | None:
    if not saved_file.exists():
        return None
    try:
        rel = saved_file.relative_to(vault_path)
    except ValueError:
        return None

    note_name = saved_file.stem + ".md"
    note_dir = saved_file.parent
    note_path = note_dir / note_name

    frontmatter = _generate_frontmatter(saved_file, note_type, tags, summary)
    file_link = "![[%s]]" % saved_file.name
    content_parts = [
        frontmatter,
        "# %s\n" % saved_file.stem.replace("_", " "),
        "## Archivo original\n",
        "Archivo: %s" % file_link,
        "Tipo: %s" % note_type,
        "Ruta: %s" % str(rel),
        "",
    ]

    if summary:
        content_parts.append("## Resumen\n")
        content_parts.append(summary)
        content_parts.append("")

    if extracted_text and len(extracted_text) > 0:
        max_chars = 10000
        text_preview = extracted_text[:max_chars]
        if len(extracted_text) > max_chars:
            text_preview += "\n\n_... (texto truncado, %d caracteres mas)_" % (
                len(extracted_text) - max_chars
            )
        content_parts.append("## Contenido extraido\n")
        content_parts.append(text_preview)

    note_path.write_text("\n".join(content_parts), encoding="utf-8")
    logger.info(
        "Companion note created: %s -> %s (%d chars extracted)",
        saved_file.name,
        note_name,
        len(extracted_text),
    )
    return note_path


def _guess_folder_from_extension(ext: str) -> str:
    return EXTENSION_FOLDERS.get(ext.lower(), "00-Inbox")


async def _classify_file_with_ai(
    filename: str, caption: str, ext: str, file_size: int, extracted_text: str = ""
) -> dict[str, Any]:
    text_sample = extracted_text[:500] if extracted_text else ""
    prompt_text = (
        "Eres un clasificador de archivos para un segundo cerebro personal (PARA + Zettelkasten). "
        "Dado el siguiente archivo, decide el nombre optimizado, carpeta de destino, "
        "tipo de nota y etiquetas.\n\n"
        "Archivo: %s\nExtension: %s\nTamano: %d bytes\n" % (filename, ext, file_size)
    )
    if caption:
        prompt_text += "Comentario del usuario: %s\n" % caption
    if text_sample:
        prompt_text += "Muestra del contenido: %s\n" % text_sample
    prompt_text += (
        "\nTipos validos: proyecto, area, recurso, nota-atomica, diario\n"
        "Carpetas: 00-Inbox, 01-Proyectos, 02-Areas/Finanzas, 02-Areas/Salud, "
        "02-Areas/Trabajo, 02-Areas/Casa, 03-Recursos, 04-Archivo\n"
        "\nResponde SOLO con JSON:\n"
        '{"name": "YYYY-MM-DD_nombre_optimizado", "folder": "03-Recursos", '
        '"type": "recurso", "tags": ["etiqueta1", "etiqueta2"], '
        '"summary": "resumen de 1 linea en espanol", "reason": "explicacion breve"}'
    )
    try:
        content = await llm.chat(
            messages=[
                {"role": "system", "content": "Eres un clasificador de archivos."},
                {"role": "user", "content": prompt_text},
            ],
            temperature=0.3,
            max_tokens=200,
        )
        import json

        match = re.search(r"\{.*\}", content, re.DOTALL)
        if match:
            result = json.loads(match.group())
            return result
    except Exception as e:
        logger.warning("AI classification failed for %s: %s", filename, e)
    return {}


async def _process_uploaded_file(
    update: Update,
    file_path: Path,
    original_name: str,
    ext: str,
    caption: str | None,
    file_size: int,
    chat_id: int,
) -> None:
    is_image = ext.lower() in IMAGE_EXTENSIONS
    if is_image:
        dest_base = ATTACHMENTS
    else:
        dest_base = INBOX

    extracted_text = ""
    if ext.lower() in EXTRACTABLE_EXTENSIONS:
        extracted_text = _extract_text_from_file(file_path)
        logger.info("Text extracted from %s: %d chars", original_name, len(extracted_text))

    ai_result = await _classify_file_with_ai(
        original_name,
        caption or "",
        ext,
        file_size,
        extracted_text,
    )
    ai_name = ai_result.get("name", "")
    ai_folder = ai_result.get("folder", "")
    ai_reason = ai_result.get("reason", "")
    ai_type = ai_result.get("type", "recurso")
    ai_tags = ai_result.get("tags", [])
    ai_summary = ai_result.get("summary", "")

    if ai_name:
        final_name = _safe_filename(ai_name)
    else:
        now_str = datetime.now().strftime("%Y-%m-%d")
        base = _safe_filename(Path(original_name).stem)[:40]
        final_name = "%s_%s" % (now_str, base)

    if ai_folder and not is_image:
        dest_dir = _safe_vault_subpath(Path("/data/obsidian_vault"), ai_folder)
    else:
        dest_dir = dest_base

    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / ("%s%s" % (final_name, ext))

    counter = 1
    while dest_path.exists():
        dest_path = dest_dir / ("%s_%d%s" % (final_name, counter, ext))
        counter += 1

    shutil.copy2(str(file_path), str(dest_path))
    size_str = _format_size(file_size)
    relative = str(dest_path.relative_to(Path("/data/obsidian_vault")))

    vault_path = Path("/data/obsidian_vault")
    note_path = None
    if extracted_text:
        note_path = _create_companion_note(
            vault_path,
            dest_path,
            extracted_text,
            ai_type,
            ai_tags,
            ai_summary,
        )

    lines = [
        "✅ Archivo clasificado y guardado:",
        "   • Nombre: %s" % dest_path.name,
        "   • Carpeta: %s" % str(dest_path.parent.relative_to(vault_path)),
        "   • Tamano: %s" % size_str,
        "   • Tipo: %s" % ai_type,
    ]
    if ai_tags:
        lines.append("   • Tags: %s" % ", ".join(ai_tags[:5]))
    if ai_reason:
        lines.append("   • Clasificacion IA: %s" % ai_reason)
    if note_path:
        lines.append("   • Nota indexada: %s ✅" % note_path.name)
    elif not is_image and ext.lower() not in EXTRACTABLE_EXTENSIONS:
        lines.append("   • ⚠️ No se pudo extraer texto (formato no soportado)")

    message_text = "\n".join(lines)
    await update.effective_message.reply_text(message_text)

    prefs = await db.get_or_create_preferences(chat_id)
    if prefs.get("voice_replies", False):
        try:
            audio_path = await text_to_speech(
                message_text.replace("✅", "Archivo clasificado y guardado.")
            )
            if audio_path:
                ogg_path = await convert_to_ogg(audio_path)
                if ogg_path:
                    with open(ogg_path, "rb") as audio_f:
                        await update.effective_message.reply_audio(audio=audio_f)
        except Exception as e:
            logger.debug("Voice reply failed for file ingest: %s", e)

    logger.info("File ingested: %s -> %s (AI: %s)", original_name, relative, ai_reason or "default")


def _format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return "%d B" % size_bytes
    elif size_bytes < 1024 * 1024:
        return "%.1f KB" % (size_bytes / 1024)
    elif size_bytes < 1024 * 1024 * 1024:
        return "%.1f MB" % (size_bytes / (1024 * 1024))
    return "%.1f GB" % (size_bytes / (1024 * 1024 * 1024))


CREDENTIALS_DIR = Path("/workspace/credentials")
CREDENTIALS_FILENAMES = {"credentials.json", "service_account.json"}


async def _handle_credentials_file(
    update: Update, file_path: Path, filename: str, user_id: int
) -> bool:
    if filename.lower() not in CREDENTIALS_FILENAMES:
        return False
    admin_ids = settings.admin_ids
    if admin_ids and user_id not in admin_ids:
        return False
    CREDENTIALS_DIR.mkdir(parents=True, exist_ok=True)
    dest = CREDENTIALS_DIR / filename
    import shutil

    shutil.copy2(str(file_path), str(dest))
    await update.effective_message.reply_text(
        "✅ *Credencial guardada exitosamente!*"
        "\n\nArchivo: `%s`"
        "\nUbicacion: `/workspace/credentials/`"
        "\n\nEjecuta `/setup_google` para verificar el estado."
        "\nLuego reinicia el bot o usa `/status` para confirmar." % filename,
        parse_mode="Markdown",
    )
    logger.info("Credentials file saved: %s (user %d)", filename, user_id)
    return True


async def document_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    user = update.effective_user
    if not message or not user:
        return
    doc: Document = message.document
    if not doc:
        return
    if doc.file_size and doc.file_size > MAX_FILE_SIZE:
        await message.reply_text("El archivo es demasiado grande (max 50 MB).")
        return
    await message.reply_chat_action("typing")
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    original_name = doc.file_name or "documento"
    ext = Path(original_name).suffix.lower()
    safe_name = _safe_filename(Path(original_name).stem)[:40]
    file_path = TEMP_DIR / ("%d_%s%s" % (user.id, safe_name, ext))
    try:
        file = await doc.get_file()
        await file.download_to_drive(file_path)
        if await _handle_credentials_file(update, file_path, original_name, user.id):
            return
        await message.reply_text("Descargando y clasificando archivo...")
        await _process_uploaded_file(
            update,
            file_path,
            original_name,
            ext,
            message.caption,
            doc.file_size or 0,
            user.id,
        )
    except Exception as e:
        logger.exception("Document handler error for user %d", user.id)
        await message.reply_text("Error al procesar el archivo: %s" % e)
    finally:
        if file_path.exists():
            file_path.unlink()


async def photo_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    user = update.effective_user
    if not message or not user:
        return
    photos = message.photo
    if not photos:
        return
    photo: PhotoSize = photos[-1]
    if photo.file_size and photo.file_size > MAX_FILE_SIZE:
        await message.reply_text("La imagen es demasiado grande (max 50 MB).")
        return
    await message.reply_chat_action("typing")
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = "photo_%s.jpg" % timestamp
    file_path = TEMP_DIR / ("%d_%s" % (user.id, file_name))
    try:
        file = await photo.get_file()
        await file.download_to_drive(file_path)
        caption = message.caption
        if caption and caption.strip():
            await message.reply_text("Analizando imagen con IA...")
            await _process_vision_image(update, user.id, str(file_path), caption, context)
        else:
            await message.reply_text("Descargando y clasificando imagen...")
            await _process_uploaded_file(
                update,
                file_path,
                file_name,
                ".jpg",
                None,
                photo.file_size or 0,
                user.id,
            )
    except Exception as e:
        logger.exception("Photo handler error for user %d", user.id)
        await message.reply_text("Error al procesar la imagen: %s" % e)
    finally:
        if file_path.exists():
            file_path.unlink()


_SEMANTIC_STOP_WORDS = {
    "guarda",
    "guardame",
    "guárdame",
    "guardar",
    "imagen",
    "foto",
    "picture",
    "siguiente",
    "esta",
    "ese",
    "esa",
    "del",
    "de",
    "la",
    "el",
    "las",
    "los",
    "un",
    "una",
    "unos",
    "unas",
    "y",
    "o",
    "a",
    "en",
    "por",
    "para",
    "con",
    "sin",
    "sobre",
    "tras",
    "me",
    "te",
    "se",
    "le",
    "les",
    "lo",
    "mi",
    "tu",
    "su",
    "es",
    "son",
    "ser",
    "estar",
    "este",
    "estos",
    "estas",
    "que",
    "como",
    "cual",
    "cuales",
    "cuando",
    "donde",
    "quien",
    "quiero",
    "puedes",
    "puede",
    "haz",
    "hazme",
    "dame",
    "muestra",
    "ver",
    "the",
    "an",
    "of",
    "in",
    "on",
    "at",
    "to",
    "for",
    "with",
    "analiza",
    "describe",
    "mira",
    "observa",
    "chequea",
    "revisa",
    "porfavor",
    "favor",
    "esto",
    "dela",
    "escudo",
    "logo",
    "simbolo",
    "captura",
}


def _extract_semantic_name(caption: str, extracted_text: str = "") -> str:
    source = (extracted_text or "").strip()
    if not source:
        source = (caption or "").strip()
    if not source:
        return "imagen_%s" % datetime.now().strftime("%Y%m%d_%H%M%S")

    patterns = [
        r"(?:escudo|logo|símbolo)\s+(?:del?\s+)?([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ]?[a-záéíóúñ]+)*)",
        r"(?:club|equipo|asociación)\s+([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ]?[a-záéíóúñ]+)*)",
        r"([A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,})*)",
    ]
    for pattern in patterns:
        matches = re.findall(pattern, source)
        if matches:
            name = matches[0].strip()
            name = re.sub(r"\s+", "_", name)
            name = re.sub(r"[^\w]", "", name)
            if len(name) > 2:
                timestamp = datetime.now().strftime("%Y%m%d")
                return "%s_%s" % (name, timestamp)

    words = re.findall(r"[a-zA-ZáéíóúÁÉÍÓÚñÑ]+", source)
    significant = [w for w in words if w.lower() not in _SEMANTIC_STOP_WORDS and len(w) > 2]
    if not significant:
        significant = [w for w in words if len(w) > 2]
    if not significant:
        return "imagen_%s" % datetime.now().strftime("%Y%m%d_%H%M%S")
    name_parts = significant[:3]
    semantic = "_".join(name_parts)
    semantic = re.sub(r"[^\w]", "", semantic)
    timestamp = datetime.now().strftime("%Y%m%d")
    return "%s_%s" % (semantic, timestamp)


async def _process_vision_image(
    update: Update, chat_id: int, image_path: str, caption: str, context=None
) -> None:
    from src.handlers.chat import TOOLS_DEFINITIONS, _execute_tool
    from src.models.schemas import MessageRole
    from src.utils.obsidian_manager import create_note_with_image, save_attachment

    if context is not None:
        context.user_data["processing_image"] = True

    await db.save_chat_message(
        chat_id, MessageRole.user.value, "[Envió una imagen. Caption: %s]" % caption[:200]
    )

    await update.effective_message.reply_text(
        "🔍 Analizando imagen con llava:7b (visión de alta precisión)..."
    )

    vision_prompt = (
        "Analiza la imagen adjunta y extrae toda la información relevante en texto plano. "
        "Si es un ticket o factura, extrae: montos, productos, fecha, establecimiento. "
        "Si es un documento, extrae el texto clave. "
        "Si es una pizarra o apunte, transcribe el contenido. "
        "Si es una captura de pantalla, describe lo importante. "
        "Si es un escudo, logo o imagen simbólica, describe qué representa con precisión: "
        "identifica colores, texto visible, animales u objetos heráldicos, y la entidad a la que pertenece. "
        "Responde en español de forma estructurada y concisa."
    )

    try:
        extracted_text = await asyncio.wait_for(
            llm.chat_vision(
                messages=[
                    {"role": "system", "content": vision_prompt},
                    {"role": "user", "content": caption},
                ],
                images=[image_path],
                temperature=0.3,
                max_tokens=512,
            ),
            timeout=180.0,
        )
    except TimeoutError:
        logger.warning("[VISION] Timeout 180s analizando imagen para user %d", chat_id)
        await update.effective_message.reply_text(
            "⚠️ El análisis de la imagen está tardando más de la cuenta debido a la carga local. "
            "Reintentando con prompt más corto..."
        )
        try:
            extracted_text = await asyncio.wait_for(
                llm.chat_vision(
                    messages=[
                        {"role": "system", "content": vision_prompt},
                        {"role": "user", "content": caption},
                    ],
                    images=[image_path],
                    temperature=0.3,
                    max_tokens=256,
                ),
                timeout=180.0,
            )
        except Exception:
            logger.exception("[VISION] Error tras timeout para user %d", chat_id)
            await update.effective_message.reply_text(
                "❌ El análisis de imagen superó el tiempo límite. Intenta con una imagen más pequeña."
            )
            if context is not None:
                context.user_data["processing_image"] = False
            return
    except Exception as e:
        logger.exception("Vision extraction error for user %d", chat_id)
        await update.effective_message.reply_text("Error al analizar la imagen: %s" % e)
        if context is not None:
            context.user_data["processing_image"] = False
        return

    gc.collect()

    try:
        await llm.unload_model(llm.vision_model)
    except Exception as _e:
        logger.warning("[VISION] No se pudo descargar llava explicitamente: %s", _e)

    gc.collect()

    logger.info("Vision extracted %d chars for user %d", len(extracted_text), chat_id)

    if not extracted_text or not extracted_text.strip():
        await update.effective_message.reply_text(
            "No pude extraer información relevante de la imagen."
        )
        await db.save_chat_message(
            chat_id,
            MessageRole.assistant.value,
            "No pude extraer información relevante de la imagen.",
        )
        if context is not None:
            context.user_data["processing_image"] = False
        return

    clean_name = _extract_semantic_name(caption, extracted_text)
    src_path = Path(image_path)

    attach_result = save_attachment(src_path, clean_name)
    if not attach_result.get("success"):
        logger.warning("[VISION] No se pudo guardar attachment: %s", attach_result.get("message"))
        await update.effective_message.reply_text(
            "⚠️ %s"
            % attach_result.get(
                "message", "No se pudo escribir en la carpeta Attachments de Obsidian."
            )
        )
    else:
        logger.info("[VISION] Imagen guardada en Obsidian: %s", attach_result.get("filename"))

    try:
        if src_path.exists():
            src_path.unlink()
    except Exception:
        pass

    del image_path

    if context is not None:
        context.user_data["processing_image"] = False

    await db.update_last_chat_message(
        chat_id,
        MessageRole.user.value,
        "[Envió una imagen. Descripción de la IA: %s]" % extracted_text.strip()[:300],
    )

    await update.effective_message.reply_text(
        "⚡ Procesando con Qwen 2.5 (razonamiento + acciones)..."
    )

    history = await db.get_chat_history(chat_id, 4)

    MAX_CONTENT_LEN = 500
    system_prompt = (
        "Eres Rafita, un asistente virtual personal. El usuario te ha enviado una imagen "
        "y un modelo de visión ya ha extraído el contenido en texto. "
        "Tu tarea es procesar ese texto extraído junto con el mensaje del usuario, "
        "y ejecutar las herramientas necesarias de forma automática.\n\n"
        "Por ejemplo:\n"
        "- Si el texto extraído contiene un gasto, usa save_expense.\n"
        "- Si contiene un evento o fecha, usa create_event.\n"
        "- Si el usuario pide guardar el contenido, usa manage_obsidian_note.\n"
        "- Si pide recordar algo, usa remember_fact o create_alert.\n\n"
        "Responde en español confirmando brevemente lo que detectaste y la acción realizada."
    )

    augmented_user_msg = "Contenido extraído de la imagen:\n%s\n\nInstrucción del usuario: %s" % (
        extracted_text.strip(),
        caption,
    )

    messages_for_llm = [{"role": "system", "content": system_prompt}]
    for msg in history[:-1]:
        c = msg.get("content", "")
        messages_for_llm.append(
            {
                "role": msg["role"],
                "content": c[:MAX_CONTENT_LEN] + ("..." if len(c) > MAX_CONTENT_LEN else ""),
            }
        )
    messages_for_llm.append({"role": "user", "content": augmented_user_msg})

    try:
        content, tool_calls = await asyncio.wait_for(
            llm.chat_with_tools(
                messages=messages_for_llm,
                tools=TOOLS_DEFINITIONS,
                max_tokens=512,
            ),
            timeout=120.0,
        )
    except TimeoutError:
        logger.error("[VISION] Qwen timeout 120s tras vision para user %d", chat_id)
        await update.effective_message.reply_text(
            "⚠️ La respuesta de texto está tardando demasiado debido a la carga local. "
            "La imagen fue guardada en Obsidian correctamente."
        )
        content = extracted_text.strip()
        tool_calls = None
    except Exception as e:
        logger.exception("Vision tool processing error for user %d", chat_id)
        await update.effective_message.reply_text(
            "Error al procesar la imagen con herramientas: %s" % e
        )
        if context is not None:
            context.user_data["processing_image"] = False
        return

    if tool_calls:
        results = []
        for tc in tool_calls:
            func_name = tc["function"]["name"]
            try:
                import json

                args = json.loads(tc["function"]["arguments"])
            except json.JSONDecodeError:
                args = {}
            logger.info("Vision tool call: user=%d tool=%s args=%s", chat_id, func_name, args)
            result = await _execute_tool(chat_id, func_name, args)
            results.append(result)

        parts = []
        if content and content.strip():
            parts.append(content)
        for r in results:
            parts.append(r.get("message", ""))
        text_to_save = "\n\n".join(parts)
    else:
        text_to_save = (
            content if content else "Procesé la imagen pero no se requirió ninguna acción."
        )

    note_title = "Imagen %s" % datetime.now().strftime("%Y-%m-%d %H%M")
    if attach_result.get("success"):
        note_result = create_note_with_image(
            title=note_title,
            body_text=text_to_save,
            image_filename=attach_result["filename"],
            folder="00-Inbox",
        )
        if note_result.get("success"):
            text_to_save += "\n\n📎 Guardado en Obsidian: %s" % note_result.get("filepath", "")
        else:
            logger.warning("[VISION] No se pudo crear nota: %s", note_result.get("message"))
            text_to_save += "\n\n⚠️ %s" % note_result.get(
                "message", "No se pudo crear la nota en Obsidian."
            )

    await db.save_chat_message(chat_id, MessageRole.assistant.value, text_to_save[:2000])
    await update.effective_message.reply_text(text_to_save[:4096])

    if context is not None:
        context.user_data["processing_image"] = False
